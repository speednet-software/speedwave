"""Create or edit a .docx from the office worker's element/ops DSL (python-docx).

Usage:
  python3 docx_build.py create <output.docx> <json-spec>
  python3 docx_build.py edit   <input.docx> <output.docx> <json-ops>

`json-spec` = {"elements": Element[]}; `json-ops` = Op[]. See ADR-055 for the shapes.
Output: {"ok": true, "path": "<output>"}
"""

from __future__ import annotations

import json

from script_runner import atomic_save, main, ok, fail


def _add_element(doc, el: dict) -> None:
    """Append one DSL element to the document."""
    etype = el.get("type")
    if etype == "heading":
        doc.add_heading(str(el["text"]), level=int(el["level"]))
    elif etype == "paragraph":
        p = doc.add_paragraph()
        run = p.add_run(str(el["text"]))
        run.bold = bool(el.get("bold", False))
        run.italic = bool(el.get("italic", False))
    elif etype == "table":
        header = list(el["header"])
        rows = [list(r) for r in el["rows"]]
        table = doc.add_table(rows=1 + len(rows), cols=len(header))
        table.style = "Table Grid"
        for j, h in enumerate(header):
            table.rows[0].cells[j].text = str(h)
        for i, row in enumerate(rows, start=1):
            for j in range(len(header)):
                table.rows[i].cells[j].text = str(row[j]) if j < len(row) else ""
    elif etype == "image":
        doc.add_picture(str(el["path"]))
    elif etype == "pagebreak":
        doc.add_page_break()
    else:
        fail(f"unknown element type: {etype}")


def _create(output: str, spec: dict) -> None:
    from docx import Document

    doc = Document()
    for el in spec.get("elements", []):
        _add_element(doc, el)
    atomic_save(output, doc.save)
    ok(path=output)


def _edit(src: str, output: str, ops: list) -> None:
    from docx import Document

    doc = Document(src)
    # `find` strings already replaced earlier in this batch: a later op with the same `find`
    # legitimately sees zero matches once consumed, and that is idempotent success, not a failure.
    already_replaced: set[str] = set()
    for op in ops:
        kind = op.get("op")
        if kind == "append":
            _add_element(doc, op["element"])
        elif kind == "replace_text":
            find, replace = str(op["find"]), str(op["replace"])
            matches = 0
            for para in doc.paragraphs:
                if find in para.text:
                    matches += 1
                    # Replace across the whole paragraph text (rebuild a single run).
                    new_text = para.text.replace(find, replace)
                    for r in list(para.runs):
                        r.text = ""
                    if para.runs:
                        para.runs[0].text = new_text
                    else:
                        para.add_run(new_text)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if find in cell.text:
                            matches += 1
                            cell.text = cell.text.replace(find, replace)
            if matches == 0 and find not in already_replaced:
                fail(
                    f"replace_text: '{find}' was not found in any paragraph or table cell "
                    "-- text may be split across runs (e.g. mixed formatting mid-phrase); "
                    "use readDocument first to confirm the exact text"
                )
            if matches > 0:
                already_replaced.add(find)
        elif kind == "delete_paragraph":
            idx = int(op["index"])
            paras = doc.paragraphs
            if idx < 0 or idx >= len(paras):
                fail(f"delete_paragraph index out of range: {idx}")
            el = paras[idx]._element
            el.getparent().remove(el)
        else:
            fail(f"unknown op: {kind}")
    atomic_save(output, doc.save)
    ok(path=output)


def _run(argv: list[str]) -> None:
    if not argv:
        fail("usage: docx_build.py create|edit ...")
    mode = argv[0]
    if mode == "create":
        if len(argv) != 3:
            fail("usage: docx_build.py create <output.docx> <json-spec>")
        _create(argv[1], json.loads(argv[2]))
    elif mode == "edit":
        if len(argv) != 4:
            fail("usage: docx_build.py edit <input.docx> <output.docx> <json-ops>")
        _edit(argv[1], argv[2], json.loads(argv[3]))
    else:
        fail(f"unknown mode: {mode}")


if __name__ == "__main__":
    main(_run)
