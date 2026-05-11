"""Extract a .docx to Markdown using python-docx (fallback engine for `readDocument`).

Usage: ``python3 python_docx_extract.py <input.docx>``
Output: ``{"ok": true, "markdown": "..."}``
"""

from __future__ import annotations

from markdown_utils import join_blocks, rows_to_markdown_table
from _io import main, ok, fail


def _heading_level(style_name: str) -> int | None:
    """Return the heading level for a Word style name like ``Heading 2``, else ``None``."""
    if not style_name:
        return None
    name = style_name.strip().lower()
    if name == "title":
        return 1
    if name.startswith("heading "):
        try:
            return min(6, max(1, int(name.split(" ", 1)[1])))
        except ValueError:
            return None
    return None


def _run(argv: list[str]) -> None:
    if len(argv) != 1:
        fail("usage: python_docx_extract.py <input.docx>")
    from docx import Document  # imported here so import errors become structured failures

    doc = Document(argv[0])
    blocks: list[str] = []

    # Walk the document body in order so tables and paragraphs interleave correctly.
    from docx.oxml.ns import qn

    body = doc.element.body
    para_iter = iter(doc.paragraphs)
    table_iter = iter(doc.tables)
    for child in body.iterchildren():
        if child.tag == qn("w:p"):
            para = next(para_iter, None)
            if para is None:
                continue
            text = para.text.strip()
            if not text:
                continue
            level = _heading_level(para.style.name if para.style else "")
            blocks.append(f"{'#' * level} {text}" if level else text)
        elif child.tag == qn("w:tbl"):
            table = next(table_iter, None)
            if table is None:
                continue
            rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
            blocks.append(rows_to_markdown_table(rows))

    ok(markdown=join_blocks(blocks))


if __name__ == "__main__":
    main(_run)
